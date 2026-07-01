"""AI Jobs Agent — main pipeline orchestrator.

Handles scraping, dedup, scoring, resume generation, reports, and dashboard.
"""

import json
import logging
import os
import re
import sys
import threading
import time
import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Dict, List, Optional, Tuple

from groq_api import query_groq
import naukri_scraper
import indeed_scraper
import linkedin_scraper
import auto_scorer
from prompts import (
    detect_profile, classify_job_profile, build_system_prompt,
    build_cover_letter_prompt, extract_delimited, extract_score,
    PROFILES,
)
from resume_reviewer import review_and_improve
from config import (
    CITIES, MIN_AI_SCORE, OUTPUT_DIR, JOBS_JSON, CV_FILE,
    SCORE_CACHE, PROGRESS_FILE, STATS_FILE, HEALTH_FILE,
    CLIENT_SECRET_FILE, TOKEN_FILE, SHEET_ID, SCOPES,
    GENERATION_MODEL, OUTPUT_DATE_DIR, MAX_WORKERS,
    HEALTH_CONSECUTIVE_ZERO_THRESHOLD, MAX_PAGES_PER_SEARCH,
    SEARCH_KEYWORDS, SEARCH_CITIES, SEARCH_MODES,
    TARGET_JOBS_COUNT,
)
from utils.fingerprint import Deduplicator
from utils.reports import (
    write_jobs_scraped, write_jobs_filtered,
    write_recommended_jobs, write_applied_jobs,
)
from utils.dashboard import generate_dashboard
from utils.metrics import MetricsTracker

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[
        logging.FileHandler("agent.log"),
        logging.StreamHandler(sys.stdout),
    ],
)
log = logging.getLogger("agent")

metrics = MetricsTracker()

# ── Google Sheets helpers ────────────────────────────────────────

def get_sheets_token() -> str:
    if not os.path.exists(TOKEN_FILE):
        log.error("No token.json found. Run: python3 auth_sheets.py step1")
        raise FileNotFoundError("token.json not found — run 'python3 auth_sheets.py step1' first")
    with open(TOKEN_FILE) as f:
        tok = json.load(f)
    if not tok.get("refresh_token"):
        log.error("token.json missing refresh_token. Re-run: python3 auth_sheets.py step1")
        raise ValueError("token.json missing refresh_token — run 'python3 auth_sheets.py step1' to re-authenticate")
    from datetime import datetime
    from google.oauth2.credentials import Credentials
    from google.auth.transport.requests import Request as AuthRequest
    expiry = None
    if tok.get("expiry"):
        try:
            expiry = datetime.fromisoformat(tok["expiry"]).replace(tzinfo=None)
        except (ValueError, TypeError):
            pass
    creds = Credentials(
        token=tok.get("token"),
        refresh_token=tok.get("refresh_token"),
        token_uri=tok.get("token_uri", "https://oauth2.googleapis.com/token"),
        client_id=tok.get("client_id"),
        client_secret=tok.get("client_secret"),
        scopes=tok.get("scopes", SCOPES),
        expiry=expiry,
    )
    if not creds.valid or not creds.token:
        log.info("Refreshing expired Google Sheets token...")
        try:
            creds.refresh(AuthRequest())
        except Exception as e:
            log.error("Token refresh failed: %s", e)
            log.error("Google Sheets access has expired or been revoked.")
            log.error("Re-authenticate by running: python3 auth_sheets.py step1")
            raise
        tok["token"] = creds.token
        tok["expiry"] = creds.expiry.replace(tzinfo=None).isoformat() if creds.expiry else tok.get("expiry")
        with open(TOKEN_FILE, "w") as f:
            json.dump(tok, f, indent=2)
        log.info("Token refreshed successfully")
    return creds.token


SHEET_RANGE = "Sheet1"

def _ensure_sheet_headers(token: str, max_retries: int = 3):
    import urllib.request, urllib.error
    url = f"https://sheets.googleapis.com/v4/spreadsheets/{SHEET_ID}/values/{SHEET_RANGE}!A1:G1"
    req = urllib.request.Request(url, headers={"Authorization": f"Bearer {token}"})
    try:
        resp = urllib.request.urlopen(req, timeout=30)
        existing = json.loads(resp.read().decode())
        if existing.get("values"):
            return
    except urllib.error.HTTPError as e:
        if e.code in (401, 403):
            body = e.read().decode()[:200]
            log.error("Sheet API auth error: %s %s", e.code, body)
            log.error("Your Google Sheets token is invalid or expired.")
            log.error("Re-authenticate: python3 auth_sheets.py step1")
            raise
    except Exception:
        pass

    headers = ["Job Title", "Company", "Match %", "Prep Topics", "Acceptance Chance %", "Status", "Date"]
    header_url = f"https://sheets.googleapis.com/v4/spreadsheets/{SHEET_ID}/values/{SHEET_RANGE}!A1:G1?valueInputOption=USER_ENTERED"
    body = json.dumps({"values": [headers]}).encode()
    for attempt in range(1, max_retries + 1):
        try:
            req = urllib.request.Request(header_url, data=body, headers={
                "Authorization": f"Bearer {token}", "Content-Type": "application/json",
            }, method="PUT")
            urllib.request.urlopen(req, timeout=30)
            return
        except Exception as e:
            log.warning("  Sheet header error (attempt %s/%s): %s", attempt, max_retries, e)
            time.sleep(attempt * 2)


def append_sheet_row(token: str, row: list, max_retries: int = 3) -> bool:
    import urllib.request, urllib.error
    url = f"https://sheets.googleapis.com/v4/spreadsheets/{SHEET_ID}/values/{SHEET_RANGE}!A:G:append?valueInputOption=USER_ENTERED&insertDataOption=INSERT_ROWS"
    body = json.dumps({"values": [row]}).encode()
    for attempt in range(1, max_retries + 1):
        try:
            req = urllib.request.Request(url, data=body, headers={
                "Authorization": f"Bearer {token}", "Content-Type": "application/json",
            })
            resp = urllib.request.urlopen(req, timeout=30)
            resp.read()
            return True
        except urllib.error.HTTPError as e:
            body_text = e.read().decode()[:200]
            if e.code == 403:
                log.error("  Sheet append failed — token expired/revoked. Re-auth needed: python3 auth_sheets.py step1")
                return False
            if e.code in (429, 500, 502, 503) and attempt < max_retries:
                time.sleep(attempt * 5)
                continue
            log.warning("  Sheet append error (attempt %s/%s): %s %s", attempt, max_retries, e.code, body_text)
            if attempt == max_retries:
                return False
        except urllib.error.URLError as e:
            if attempt < max_retries:
                time.sleep(attempt * 5)
                continue
            log.error("  Sheet network error: %s", e)
            return False
    return False


# ── Pipeline helpers ────────────────────────────────────────────

def load_json(path: str) -> list:
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def save_json(path: str, data: list) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def sanitize_folder_name(s: str) -> str:
    s = re.sub(r'[^\w\s-]', '', s).strip().lower()
    return re.sub(r'[-\s]+', '_', s)[:60]


def load_cv(path: str = CV_FILE) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def load_progress(path: str = PROGRESS_FILE) -> dict:
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}


def save_progress(path: str = PROGRESS_FILE, data: dict = None) -> None:
    if data is None:
        data = load_progress(path)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def _detect_broken_scraper(name: str, jobs_found: int, health_data: dict) -> dict:
    if name not in health_data:
        health_data[name] = {
            "last_run": datetime.datetime.now().isoformat(),
            "jobs_found": jobs_found,
            "http_success_rate": 1.0,
            "consecutive_zero_jobs": 0,
            "healthy": True,
        }
    h = health_data[name]
    h["last_run"] = datetime.datetime.now().isoformat()
    if jobs_found == 0:
        h["consecutive_zero_jobs"] = h.get("consecutive_zero_jobs", 0) + 1
        if h["consecutive_zero_jobs"] >= HEALTH_CONSECUTIVE_ZERO_THRESHOLD:
            h["healthy"] = False
            log.warning("  WARNING: %s scraper may be BROKEN (zero jobs for %s+ runs)", name, HEALTH_CONSECUTIVE_ZERO_THRESHOLD)
    else:
        h["consecutive_zero_jobs"] = 0
        h["healthy"] = True
    h["jobs_found"] = jobs_found
    health_data[name] = h
    return health_data


def _load_health() -> dict:
    if os.path.exists(HEALTH_FILE):
        try:
            with open(HEALTH_FILE, "r") as f:
                return json.load(f)
        except (json.JSONDecodeError, TypeError):
            pass
    return {}


def _save_health(health: dict) -> None:
    with open(HEALTH_FILE, "w") as f:
        json.dump(health, f, indent=2)


def _print_separator(title: str) -> None:
    log.info("")
    log.info("%s", "=" * 55)
    log.info("%s", title)
    log.info("%s", "=" * 55)


def _print_scraper_summary(name: str, stats: dict, health: dict) -> None:
    h = health.get(name.lower(), {})
    status = "OK" if h.get("healthy", True) else "BROKEN"
    log.info("%s:", name)
    log.info("  Searches executed:       %s", stats.get("queries", 0))
    log.info("  Pages scraped:           %s", stats.get("pages_scraped", 0))
    log.info("  Jobs found:              %s", stats.get("jobs_found", 0))
    log.info("  Duplicates removed:      %s", stats.get("duplicates", 0))
    log.info("  Failed requests:         %s", stats.get("failed_requests", 0))
    log.info("  Early stopped searches:  %s", stats.get("early_stopped", 0))
    log.info("  Duration:                %.1fs", stats.get("total_duration", 0))
    log.info("  Health status:           %s", status)


# ── Resume generation ───────────────────────────────────────────

def _resume_fingerprint(job: Dict[str, Any]) -> str:
    """Create a fingerprint for resume reuse detection."""
    title = sanitize_folder_name(job.get("title", ""))
    category = sanitize_folder_name(job.get("category", ""))
    return f"{category}_{title}"


def _check_resume_reuse(
    job: Dict[str, Any],
    progress: Dict[str, Any],
    output_base: str,
) -> Optional[str]:
    """Check if a similar resume already exists and can be reused."""
    fp = _resume_fingerprint(job)
    if fp in progress:
        folder = progress[fp].get("folder", "")
        cv_path = os.path.join(output_base, folder, "tailored_cv.docx") if folder else ""
        if cv_path and os.path.exists(cv_path):
            log.info("  Reusing existing resume (same role/category)")
            return folder
    return None


def _strip_explanatory_notes(text: str) -> str:
    """Remove any LLM-generated preamble/explanation from CV text."""
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        lowered = line.strip().lower()
        if lowered.startswith(("i've", "i have", "i added", "i've added", "note:", "note -")):
            continue
        if "i've added relevant keywords" in lowered:
            continue
        cleaned.append(line)
    return "\n".join(cleaned).strip()


def generate_for_job(job: Dict[str, Any], cv_text: str) -> Tuple[Optional[str], Optional[str], Optional[str], Optional[int], str]:
    """Generate tailored CV, cover letter, prep topics, and acceptance chance."""
    profile = classify_job_profile(job)
    log.info("  Target profile: %s", profile)

    system_prompt = build_system_prompt(profile, cv_text, job.get("description", ""))
    user_prompt = f"""TARGET JOB:
Title: {job['title']}
Company: {job['company']}
Location: {job['location']}
Category: {job.get('category', 'N/A')}

JOB DESCRIPTION:
{job.get('description', 'Not available')}

CANDIDATE'S ORIGINAL PROFILE (use ONLY this information, NEVER invent):
{cv_text}

Produce the 4 items for this role."""

    response = query_groq(system_prompt, user_prompt, model=GENERATION_MODEL)

    cv = _strip_explanatory_notes(extract_delimited(response, "TAILORED_CV"))
    cl = extract_delimited(response, "COVER_LETTER")
    prep = extract_delimited(response, "INTERVIEW_PREP")
    chance = extract_delimited(response, "ACCEPTANCE_CHANCE")

    if not cv:
        cv = response
    if not cl:
        cl_prompt = build_cover_letter_prompt(profile, job, cv_text)
        cl = query_groq(cl_prompt, f"Write a cover letter for {job['title']} at {job['company']}.", model=GENERATION_MODEL)

    chance_num = extract_score(chance) if chance else 50

    # Quality review pass
    log.info("  Running quality review...")
    try:
        cv = _strip_explanatory_notes(review_and_improve(cv, job.get("title", ""), job.get("description", "")))
    except Exception as e:
        log.warning("  Quality review skipped: %s", e)

    return cv, cl, prep, chance_num, profile


# ── DOCX/PDF generation ─────────────────────────────────────────

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF

SKILL_CATEGORIES = [
    "Data Analytics:", "Visualization & Reporting:",
    "AI / ML & GenAI:", "Other Skills:",
    "Programming:", "Agentic AI:", "SAP & ERP:",
    "Finance & Accounting:",
]

_MONTHS = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
_DATE_PATTERN = re.compile(rf"{_MONTHS}\s+\d{{4}}\s*(?:[–—-]|to|–|—)\s*{_MONTHS}\s+\d{{4}}")
_BODY_FONT = "Calibri"
_HEADING_FONT = "Calibri"
_BODY_SIZE = Pt(10.5)
_HEADING_SIZE = Pt(13)
_CONTACT_SIZE = Pt(9)
_NAVY = RGBColor(0x1F, 0x3A, 0x5F)
_DARK_GRAY = RGBColor(0x55, 0x55, 0x55)


def save_docx(text: str, path: str) -> None:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = _BODY_FONT
    style.font.size = _BODY_SIZE
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    lines = _prepare_lines(text)
    _render_docx(doc, lines)
    doc.save(path)


def save_pdf(text: str, path: str) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_font("DejaVu", "", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")
    pdf.add_font("DejaVu", "B", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf")
    pdf.set_margins(20, 18, 20)
    _render_pdf(pdf, _prepare_lines(text))
    pdf.output(path)


def _prepare_lines(text: str) -> list:
    raw = text.strip().split("\n")
    return [l.strip() for l in raw if l.strip()]


def _which_section(line: str):
    u = line.strip().upper().rstrip(":")
    section_keywords = [
        ("PROFESSIONAL SUMMARY", "summary"),
        ("TECHNICAL SKILLS", "skills"),
        ("WORK EXPERIENCE", "work"),
        ("GITHUB PROJECTS", "projects"),
        ("DSA PROJECTS", "dsa_projects"),
        ("EDUCATION", "education"),
        ("CERTIFICATIONS", "certs"),
    ]
    for kw, label in section_keywords:
        if u == kw or u.startswith(kw):
            if kw.startswith("PROFESSIONAL"):
                return "summary"
            if kw.startswith("TECHNICAL"):
                return "skills"
            if kw.startswith("WORK"):
                return "work"
            if kw.startswith("GITHUB"):
                return "projects"
            if kw.startswith("DSA"):
                return "dsa_projects"
            if kw.startswith("EDUCATION"):
                return "education"
            if kw.startswith("CERTIFICATION"):
                return "certs"
    return None


def _render_docx(doc: Document, lines: list) -> None:
    current_section = None
    for line in lines:
        sec = _which_section(line)
        if sec:
            p = doc.add_paragraph()
            run = p.add_run(line.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = _NAVY
            run.font.name = _HEADING_FONT
            current_section = sec
            continue

        if current_section == "summary":
            p = doc.add_paragraph()
            p.add_run(line).font.size = Pt(10)
            continue

        if current_section == "skills":
            if any(line.upper().startswith(cat.rstrip(":").upper()) for cat in SKILL_CATEGORIES):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
            elif line.startswith("-"):
                p = doc.add_paragraph(style='List Bullet')
                p.text = line.lstrip("- ").strip()
                p.style.font.size = Pt(9.5)
            continue

        if current_section == "work":
            if not line.startswith("-") and "—" in line:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10.5)
            elif not line.startswith("-"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
            elif line.startswith("-"):
                p = doc.add_paragraph(style='List Bullet')
                p.text = line.lstrip("- ").strip()
            continue

        if current_section in ("projects", "dsa_projects"):
            if not line.startswith("-") and not line.startswith("GitHub") and not line.startswith("GreatLearning"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(10)
            elif line.startswith("GitHub") or line.startswith("GreatLearning"):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(9)
                run.font.color.rgb = _NAVY
            elif line.startswith("-"):
                p = doc.add_paragraph(style='List Bullet')
                p.text = line.lstrip("- ").strip()
            continue

        if line.startswith("-"):
            p = doc.add_paragraph(style='List Bullet')
            p.text = line.lstrip("- ").strip()
        else:
            p = doc.add_paragraph()
            p.add_run(line).font.size = Pt(10)


def _render_pdf(pdf: FPDF, lines: list) -> None:
    lm = pdf.l_margin
    avail_w = pdf.w - lm - pdf.r_margin
    current_section = None

    for line in lines:
        sec = _which_section(line)
        if sec:
            pdf.set_font("DejaVu", "B", 11)
            pdf.set_text_color(0x1F, 0x3A, 0x5F)
            pdf.multi_cell(avail_w, 7, line.upper(), align="L")
            pdf.set_text_color(0, 0, 0)
            current_section = sec
            continue

        if current_section == "summary":
            pdf.set_font("DejaVu", "", 9.5)
            pdf.multi_cell(avail_w, 5.5, line)
            continue

        if current_section == "skills":
            if line.startswith("-"):
                pdf.set_font("DejaVu", "", 9)
                pdf.set_x(lm + 6)
                pdf.multi_cell(avail_w - 6, 5, "- " + line.lstrip("- ").strip())
            else:
                pdf.set_font("DejaVu", "B", 10)
                pdf.set_x(lm)
                pdf.multi_cell(avail_w, 5.5, line)
            continue

        if current_section == "work":
            if not line.startswith("-") and "—" in line:
                pdf.set_font("DejaVu", "B", 10.5)
                pdf.set_x(lm)
                pdf.multi_cell(avail_w, 6, line)
            elif not line.startswith("-"):
                pdf.set_font("DejaVu", "B", 10)
                pdf.set_x(lm)
                pdf.multi_cell(avail_w, 5.5, line)
            elif line.startswith("-"):
                pdf.set_font("DejaVu", "", 9)
                pdf.set_x(lm + 6)
                pdf.multi_cell(avail_w - 6, 5, "- " + line.lstrip("- ").strip())
            continue

        if current_section in ("projects", "dsa_projects"):
            if line.startswith("GitHub") or line.startswith("GreatLearning"):
                pdf.set_font("DejaVu", "", 8.5)
                pdf.set_text_color(0x1F, 0x3A, 0x5F)
                pdf.set_x(lm)
                pdf.multi_cell(avail_w, 4.5, line)
                pdf.set_text_color(0, 0, 0)
            elif line.startswith("-"):
                pdf.set_font("DejaVu", "", 9)
                pdf.set_x(lm + 6)
                pdf.multi_cell(avail_w - 6, 5, "- " + line.lstrip("- ").strip())
            else:
                pdf.set_font("DejaVu", "B", 10)
                pdf.set_x(lm)
                pdf.multi_cell(avail_w, 5.5, line)
            continue

        if line.startswith("-"):
            pdf.set_font("DejaVu", "", 9)
            pdf.set_x(lm + 6)
            pdf.multi_cell(avail_w - 6, 5, "- " + line.lstrip("- ").strip())
        else:
            pdf.set_font("DejaVu", "", 9.5)
            pdf.set_x(lm)
            pdf.multi_cell(avail_w, 5.5, line)


# ── Scraping phase ─────────────────────────────────────────────

def run_scraper_internal(scraper_module, name: str, job_age_hours: int, stop_event: threading.Event = None) -> Tuple[List[Dict], int, Dict]:
    """Run a scraper module and merge results."""
    total_jobs = []
    total_new = 0
    scraper_stats = {}

    try:
        if hasattr(scraper_module, 'fetch_all') and callable(scraper_module.fetch_all):
            jobs = scraper_module.fetch_all(job_age_hours=job_age_hours, stop_event=stop_event)
        else:
            jobs = scraper_module.fetch_all()

        if jobs:
            added = scraper_module.merge_into_all_roles(jobs)
            total_jobs = jobs
            total_new = added
        else:
            total_jobs = []
            total_new = 0

        if hasattr(scraper_module, 'get_last_stats'):
            scraper_stats = scraper_module.get_last_stats()
    except Exception as e:
        log.error("  %s scraper failed: %s", name, e)
        return [], 0, scraper_stats

    return total_jobs, total_new, scraper_stats


def run_scraping_phase(
    parallel: bool, job_age_hours: int,
    target_jobs: int = 0, cv_text: str = "",
) -> Tuple[List[Dict], Dict, Dict, List[Dict]]:
    """Run configured scrapers, score in real-time, stop when target_jobs high-scoring jobs collected.

    Returns (all_fetched_jobs, source_stats, health_data, good_jobs).
    When target_jobs <= 0, scoring is skipped and good_jobs is empty.
    """
    _print_separator("SCRAPING PHASE")
    metrics.start_phase("scraping")

    stop_event = threading.Event()
    all_jobs = []
    good_jobs: List[Dict] = []
    good_jobs_lock = threading.Lock()
    source_stats = {}
    health_data = _load_health()

    scraper_configs = []
    from config import WEBSITES
    if WEBSITES.get("naukri", True):
        scraper_configs.append((naukri_scraper, "Naukri"))
    if WEBSITES.get("indeed", True):
        scraper_configs.append((indeed_scraper, "Indeed"))
    if WEBSITES.get("linkedin", True):
        scraper_configs.append((linkedin_scraper, "LinkedIn"))

    if not scraper_configs:
        log.warning("No scrapers enabled in search_config.json")
        return [], source_stats, health_data, []

    log.info("Job age filter: %d hours", job_age_hours)
    if target_jobs > 0:
        log.info("Target: %s high-scoring jobs (score >= %s)", target_jobs, MIN_AI_SCORE)

    # Load already-scored keys to avoid duplicate scoring
    existing_scores = load_json(SCORE_CACHE)
    already_scored: set = set()
    for s in existing_scores:
        already_scored.add((s.get("title", ""), s.get("company", "")))

    def _score_and_collect(jobs_list: List[Dict]) -> None:
        """Score a batch of new jobs and collect those above threshold."""
        for job in jobs_list:
            if stop_event.is_set():
                return
            key = (job.get("title", ""), job.get("company", ""))
            if key in already_scored:
                continue
            already_scored.add(key)
            try:
                entry = auto_scorer.score_single_job_to_entry(job, cv_text)
                if entry is None:
                    continue
                auto_scorer.append_single_score(entry)
                if entry["score"] >= MIN_AI_SCORE:
                    with good_jobs_lock:
                        good_jobs.append(entry)
                        if len(good_jobs) >= target_jobs:
                            log.info(
                                "Target of %s high-scoring jobs reached! Stopping scrapers.",
                                target_jobs,
                            )
                            stop_event.set()
                            return
            except Exception as e:
                log.warning("  Score failed for %s @ %s: %s",
                            job.get("title"), job.get("company"), e)

    if parallel and len(scraper_configs) > 1:
        log.info("Running scrapers in parallel mode...")
        with ThreadPoolExecutor(max_workers=len(scraper_configs)) as executor:
            futures = {}
            for mod, name in scraper_configs:
                future = executor.submit(
                    run_scraper_internal, mod, name, job_age_hours, stop_event,
                )
                futures[future] = (mod, name)

            for future in as_completed(futures):
                if stop_event.is_set():
                    for f in futures:
                        if not f.done():
                            f.cancel()
                    break
                mod, name = futures[future]
                try:
                    jobs, added, stats = future.result()
                    log.info("%s: %s jobs fetched, %s new", name, len(jobs), added)
                    source_stats[name.lower()] = stats
                    health_data = _detect_broken_scraper(name, len(jobs), health_data)
                    all_jobs.extend(jobs)
                    if target_jobs > 0 and cv_text:
                        _score_and_collect(jobs)
                except Exception as e:
                    log.error("  %s scraper failed: %s", name, e)
                    source_stats[name.lower()] = {}
                    health_data = _detect_broken_scraper(name, 0, health_data)
    else:
        for mod, name in scraper_configs:
            if stop_event.is_set():
                log.info("%s: skipped (target reached)", name)
                continue
            log.info("Fetching %s jobs...", name)
            jobs, added, stats = run_scraper_internal(mod, name, job_age_hours, stop_event)
            log.info("%s: %s jobs fetched, %s new", name, len(jobs), added)
            source_stats[name.lower()] = stats
            health_data = _detect_broken_scraper(name, len(jobs), health_data)
            all_jobs.extend(jobs)
            if target_jobs > 0 and cv_text:
                _score_and_collect(jobs)

    _save_health(health_data)
    metrics.end_phase("scraping")

    good_jobs.sort(key=lambda x: x["score"], reverse=True)
    if target_jobs > 0:
        log.info("Scraping complete: %s total jobs, %s high-scoring (score >= %s)",
                 len(all_jobs), len(good_jobs), MIN_AI_SCORE)

    return all_jobs, source_stats, health_data, good_jobs[:target_jobs] if target_jobs > 0 else []


def log_scraper_results(name: str, jobs: list, added: int) -> None:
    log.info("%s fetched %s jobs, %s new", name, len(jobs), added)


# ── Main ────────────────────────────────────────────────────────

def main():
    metrics.start_phase("total")

    # Parse CLI args
    parallel = "--parallel" in sys.argv
    only_mode = "--only" in sys.argv
    skip_scrape = "--skip-scrape" in sys.argv
    skip_score = "--skip-score" in sys.argv
    skip_generate = "--skip-generate" in sys.argv
    max_jobs = None
    gen_workers = 3
    target_jobs = TARGET_JOBS_COUNT
    for i, arg in enumerate(sys.argv):
        if arg == "--max-jobs" and i + 1 < len(sys.argv):
            try:
                max_jobs = int(sys.argv[i + 1])
            except ValueError:
                pass
        elif arg == "--workers" and i + 1 < len(sys.argv):
            try:
                gen_workers = int(sys.argv[i + 1])
            except ValueError:
                pass
        elif arg == "--target-jobs" and i + 1 < len(sys.argv):
            try:
                target_jobs = int(sys.argv[i + 1])
            except ValueError:
                pass

    job_age_hours = 24

    all_fetched_jobs = []
    source_stats = {}
    health_data = {}
    scored_during_scrape = False
    good = []

    # ── Phase 1: Scrape + optional real-time scoring ────────────
    cv_text = ""  # loaded before scrape if scoring during scrape
    if not skip_scrape:
        do_realtime_scoring = not skip_score and target_jobs > 0
        if do_realtime_scoring:
            cv_text = load_cv(CV_FILE)
            log.info("Real-time scoring enabled: target %s high-scoring jobs", target_jobs)

        all_fetched_jobs, source_stats, health_data, good = run_scraping_phase(
            parallel, job_age_hours,
            target_jobs=target_jobs if do_realtime_scoring else 0,
            cv_text=cv_text if do_realtime_scoring else "",
        )
        scored_during_scrape = bool(good)

        # Cross-platform dedup
        metrics.start_phase("dedup")
        dedup = Deduplicator()
        unique_jobs = []
        dup_count = 0
        for job in all_fetched_jobs:
            if dedup.is_duplicate(
                job.get("company", ""),
                job.get("title", ""),
                job.get("location", ""),
                job.get("url", ""),
            ):
                dup_count += 1
            else:
                unique_jobs.append(job)
        all_fetched_jobs = unique_jobs
        log.info("Cross-platform dedup: %s duplicates removed, %s unique jobs", dup_count, len(unique_jobs))
        metrics.end_phase("dedup")
        metrics.increment("total_fetched", len(all_fetched_jobs))
        metrics.increment("duplicates_removed", dup_count)
    else:
        log.info("Skipping scrape phase (--skip-scrape)")
        all_fetched_jobs = load_json(JOBS_JSON)

    # ── Phase 2: Filter blacklisted ─────────────────────────────
    metrics.start_phase("filter")
    from config import BLACKLIST
    filtered_jobs = []
    rejected_jobs = []
    blacklist_lower = [b.lower() for b in BLACKLIST]
    for job in all_fetched_jobs:
        company = (job.get("company") or "").lower()
        title = (job.get("title") or "").lower()
        is_blacklisted = any(b in company or b in title for b in blacklist_lower)
        if is_blacklisted:
            job["rejection_reason"] = "Blacklisted company/keyword"
            rejected_jobs.append(job)
        else:
            filtered_jobs.append(job)
    all_jobs = filtered_jobs
    log.info("Blacklist filter: %s rejected, %s kept", len(rejected_jobs), len(all_jobs))
    metrics.end_phase("filter")

    # ── Phase 3: Score (only if not already scored during scrape) ──
    if not skip_score and not scored_during_scrape:
        metrics.start_phase("scoring")
        _print_separator("SCORING PHASE")
        log.info("Running weighted auto-scorer...")
        cv_text = load_cv(CV_FILE)
        try:
            added_scores = auto_scorer.score_all_unscored()
            log.info("Scored %s new jobs", added_scores)
            metrics.increment("jobs_scored", added_scores)
        except Exception as e:
            log.error("Auto-scorer failed: %s", e)
        metrics.end_phase("scoring")

    # ── Phase 4: Load scores & get recommended ──────────────────
    if not good:
        if not cv_text:
            cv_text = load_cv(CV_FILE)
        scored = load_json(SCORE_CACHE)

        if skip_score:
            for job in all_jobs:
                good.append({
                    "title": job.get("title", ""),
                    "company": job.get("company", ""),
                    "category": job.get("category", ""),
                    "score": MIN_AI_SCORE,
                    "location": job.get("location", ""),
                })
        else:
            good = [r for r in scored if isinstance(r.get('score'), (int, float)) and r['score'] >= MIN_AI_SCORE]
            good.sort(key=lambda x: x['score'], reverse=True)

    metrics.increment("total_scored", len(load_json(SCORE_CACHE)))
    metrics.increment("recommended", len(good))

    _print_separator("RECOMMENDED JOBS")
    log.info("Jobs with score >= %s: %s total", MIN_AI_SCORE, len(good))
    log.info("%-3s %-6s %-18s %-40s %-22s", "#", "Score", "Category", "Job Title", "Company")
    log.info("%s", "-" * 90)
    for i, r in enumerate(good, 1):
        cat = r.get('category', 'N/A')[:16]
        title = r['title'][:38]
        company = r['company'][:20]
        log.info("%-3s %-4s/10 %-18s %-38s %-20s", i, r['score'], cat, title, company)
    log.info("%s", "-" * 90)

    if not good:
        log.info("No jobs above threshold. Exiting.")
        return

    # ── Phase 5: Generate resumes ───────────────────────────────
    completed = 0
    failed = 0
    skipped = 0
    reused = 0
    sheet_uploads = 0
    sheet_failures = 0
    applied_jobs_report = []

    if not skip_generate:
        _print_separator("GENERATION PHASE")

        log.info("Connecting to Google Sheets...")
        sheets_token = None
        try:
            sheets_token = get_sheets_token()
            _ensure_sheet_headers(sheets_token)
            log.info("Google Sheets ready")
        except Exception as e:
            log.warning("  Google Sheets auth/init failed — jobs will not be uploaded to sheet: %s", e)

        os.makedirs(OUTPUT_DIR, exist_ok=True)
        os.makedirs(OUTPUT_DATE_DIR, exist_ok=True)

        all_jobs_map = {}
        for job in load_json(JOBS_JSON):
            key = (job.get("title", ""), job.get("company", ""))
            all_jobs_map[key] = job

        progress = load_progress(PROGRESS_FILE)

        # Filter out already-generated jobs so we always get N new ones
        good_ungenerated = []
        for r in good:
            job_key = sanitize_folder_name(f"{r['company']}_{r['title']}")
            if job_key not in progress or progress[job_key].get("status") != "done":
                good_ungenerated.append(r)
        skipped_old = len(good) - len(good_ungenerated)
        if skipped_old:
            log.info("Skipped %s already-generated jobs, %s remaining", skipped_old, len(good_ungenerated))
        good = good_ungenerated

        if max_jobs and max_jobs < len(good):
            log.info("Limiting to %s jobs (--max-jobs)", max_jobs)
            good = good[:max_jobs]

        def _gen_worker(job, r, cv_text):
            """Thread worker: generate files, return result dict or None."""
            log.info("  (Score: %s/10) %s @ %s", r['score'], r['title'], r['company'])
            cv_result, cl_result, prep_result, chance_num, profile = generate_for_job(job, cv_text)
            if not cv_result:
                return None
            from docx import Document as DocxDoc
            folder_name = sanitize_folder_name(f"{r['company']}_{r['title']}")
            folder_path = os.path.join(OUTPUT_DATE_DIR, folder_name)
            os.makedirs(folder_path, exist_ok=True)
            save_docx(cv_result, os.path.join(folder_path, "tailored_cv.docx"))
            save_pdf(cv_result, os.path.join(folder_path, "tailored_cv.pdf"))
            if cl_result:
                save_docx(cl_result, os.path.join(folder_path, "cover_letter.docx"))
                save_pdf(cl_result, os.path.join(folder_path, "cover_letter.pdf"))
            return {
                "folder_name": folder_name, "prep_result": prep_result,
                "chance_num": chance_num, "profile": profile,
            }

        log.info("Generating with %s worker(s)...", gen_workers)
        total = len(good)
        with ThreadPoolExecutor(max_workers=gen_workers) as executor:
            futures = {}
            for i, r in enumerate(good, 1):
                key = (r["title"], r["company"])
                job = all_jobs_map.get(key)
                if not job:
                    log.warning("[%s/%s] SKIP (no data): %s @ %s", i, total, r['title'], r['company'])
                    skipped += 1
                    continue

                job_key = sanitize_folder_name(f"{r['company']}_{r['title']}")
                if job_key in progress:
                    log.info("[%s/%s] SKIP (already done): %s @ %s", i, total, r['title'], r['company'])
                    skipped += 1
                    continue

                reuse = _check_resume_reuse(job, progress, OUTPUT_DIR)
                if reuse:
                    reused += 1
                    progress[job_key] = {"status": "reused", "folder": reuse}
                    save_progress(PROGRESS_FILE, progress)
                    continue

                log.info("[%s/%s] Submitting: (Score: %s/10) %s @ %s", i, total, r['score'], r['title'], r['company'])
                future = executor.submit(_gen_worker, job, r, cv_text)
                futures[future] = (i, r, job, job_key)

            for future in as_completed(futures):
                i, r, job, job_key = futures[future]
                metrics.start_phase(f"gen_{i}")
                try:
                    result = future.result()
                    if result:
                        completed += 1
                        folder_name = result["folder_name"]
                        prep_result = result["prep_result"]
                        chance_num = result["chance_num"]
                        profile = result["profile"]

                        if sheets_token:
                            match_pct = r['score'] * 10
                            today = datetime.date.today().strftime("%Y-%m-%d")
                            row = [
                                r['title'], r['company'], match_pct,
                                prep_result or "", chance_num or 50, "Applied", today,
                            ]
                            try:
                                if append_sheet_row(sheets_token, row):
                                    sheet_uploads += 1
                                else:
                                    sheet_failures += 1
                            except Exception as e:
                                log.error("  Sheet append failed: %s @ %s: %s", r['title'], r['company'], e)
                                sheet_failures += 1

                        applied_jobs_report.append({
                            "title": r['title'],
                            "company": r['company'],
                            "location": job.get("location", ""),
                            "score": r['score'],
                            "profile": profile,
                            "match_pct": r['score'] * 10,
                            "folder": folder_name,
                        })
                        progress[job_key] = {"status": "done", "folder": folder_name}
                    else:
                        log.warning("[%s/%s] No CV generated: %s @ %s", i, total, r['title'], r['company'])

                except Exception as e:
                    log.error("  FAILED [%s/%s] %s @ %s: %s", i, total, r['title'], r['company'], e)
                    failed += 1

                metrics.end_phase(f"gen_{i}")
                save_progress(PROGRESS_FILE, progress)

        log.info("Generation complete: %s completed, %s reused, %s skipped, %s failed",
                 completed, reused, skipped, failed)
        log.info("Sheet updates: %s uploaded, %s failed", sheet_uploads, sheet_failures)

        write_applied_jobs(applied_jobs_report)

    # ── Phase 6: Reports ────────────────────────────────────────
    _print_separator("REPORTS")
    write_jobs_scraped(all_fetched_jobs)
    write_jobs_filtered(rejected_jobs)
    write_recommended_jobs(good)

    # Compute daily trend
    daily_trend = {}
    for job in all_fetched_jobs:
        d = job.get("date", datetime.date.today().isoformat())
        daily_trend[d] = daily_trend.get(d, 0) + 1

    # Top companies
    company_counts = {}
    for job in all_fetched_jobs:
        c = job.get("company", "Unknown")
        company_counts[c] = company_counts.get(c, 0) + 1
    top_companies = sorted(
        [{"name": k, "count": v} for k, v in company_counts.items()],
        key=lambda x: x["count"], reverse=True
    )[:15]

    # Top categories
    cat_counts = {}
    for job in all_fetched_jobs:
        cat = job.get("category", "Unknown")
        cat_counts[cat] = cat_counts.get(cat, 0) + 1
    top_categories = sorted(
        [{"name": k, "count": v} for k, v in cat_counts.items()],
        key=lambda x: x["count"], reverse=True
    )

    dashboard_stats = {
        "date": datetime.date.today().isoformat(),
        "runtime_seconds": round(metrics.total_elapsed(), 1),
        "total_scraped": len(all_fetched_jobs),
        "total_recommended": len(good),
        "total_applied": len([j for j in load_json(PROGRESS_FILE).values() if j.get("status") == "done"]),
        "source_stats": source_stats,
        "daily_trend": daily_trend,
        "top_companies": top_companies,
        "top_categories": top_categories,
    }
    generate_dashboard(dashboard_stats)

    # ── Phase 7: Summary ────────────────────────────────────────
    _print_separator("SCRAPING SUMMARY")
    for name_key in ["naukri", "indeed", "linkedin"]:
        if name_key in source_stats:
            _print_scraper_summary(name_key.capitalize(), source_stats.get(name_key, {}), health_data)
        else:
            log.info("%s: (not fetched)", name_key.capitalize())

    _print_separator("TOTALS")
    all_jobs_count = len(load_json(JOBS_JSON))
    log.info("Jobs in database:         %s", all_jobs_count)
    scored_count = len(load_json(SCORE_CACHE))
    log.info("Jobs scored:              %s", scored_count)
    if good:
        avg = sum(r['score'] for r in good) / len(good)
        log.info("Average score (good):     %.1f/10", avg)
    log.info("Jobs above threshold:     %s", len(good))
    log.info("Applications generated:   %s", completed)

    metrics.print_summary()

    # Save agent stats
    stats_entry = {
        "date": datetime.date.today().isoformat(),
        "runtime_seconds": round(metrics.total_elapsed(), 1),
        "total_scraped": len(all_fetched_jobs),
        "total_scored": scored_count,
        "recommended": len(good),
        "sources": source_stats,
    }
    existing_stats = load_json(STATS_FILE)
    if not isinstance(existing_stats, list):
        existing_stats = [existing_stats] if existing_stats else []
    existing_stats.append(stats_entry)
    save_json(STATS_FILE, existing_stats)

    _print_separator("DONE")
    log.info("All reports in outputs/reports/")
    log.info("Dashboard: outputs/dashboard/index.html")
    log.info("Generated files: %s/", OUTPUT_DATE_DIR)


if __name__ == "__main__":
    main()
